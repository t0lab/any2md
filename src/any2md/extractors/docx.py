"""docx extractor — python-docx based, pure-Python."""
from __future__ import annotations

import logging
import posixpath
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from any2md.extractors._image_utils import image_dimensions, normalize_to_png
from any2md.extractors._smartart import (
    detect_layout_kind,
    parse_rels,
    parse_smartart_tree,
)
from any2md.extractors.base import ExtractResult

log = logging.getLogger(__name__)

_HEADING_RE = re.compile(r"heading\s*(\d+)", re.IGNORECASE)
# Styles whose name carries no signal worth surfacing on a ParagraphBlock.
_DEFAULT_STYLES = {"", "Normal", "Body Text", "Default Paragraph Font", "List Paragraph"}
_CODE_STYLES = {"HTML Preformatted", "Preformatted Text", "Macro Text"}
_QUOTE_STYLES = {"Quote", "Intense Quote"}
# These namespaces are not in python-docx's nsmap; reference the URIs directly.
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_DGM_NS = "http://schemas.openxmlformats.org/drawingml/2006/diagram"


class DocxExtractor:
    format = "docx"

    def extract(self, source: Path) -> ExtractResult:
        warnings: list[str] = []
        images: dict[str, bytes] = {}
        counter = {"img": 0}

        doc = Document(str(source))
        numbering = self._build_numbering_map(doc, warnings)
        self._load_diagrams(source, warnings)

        blocks: list[dict[str, Any]] = []
        for item in self._iter_block_items(doc.element.body, doc):
            if isinstance(item, Paragraph):
                new_blocks = self._paragraph_to_blocks(
                    item, doc, numbering, images, counter, warnings
                )
            else:  # Table
                new_blocks = [self._table_to_block(item)]
            for b in new_blocks:
                b["order"] = len(blocks)
                blocks.append(b)

        self._wire_image_neighbors(blocks)

        ir: dict[str, Any] = {
            "format": "docx",
            "source": str(source),
            "blocks": blocks,
        }
        return ExtractResult(ir=ir, images=images, warnings=warnings)

    # ----------------- body walk -----------------

    def _iter_block_items(self, parent_elm: Any, parent: Any):
        """Yield Paragraph / Table in document order.

        `doc.paragraphs` and `doc.tables` are separate lists — iterating the
        body element directly is the only way to preserve their interleaving.
        """
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    # ----------------- paragraphs -----------------

    def _paragraph_to_blocks(
        self,
        para: Paragraph,
        doc: Any,
        numbering: dict[str, dict[int, str]],
        images: dict[str, bytes],
        counter: dict[str, int],
        warnings: list[str],
    ) -> list[dict[str, Any]]:
        style_name = ""
        try:
            if para.style is not None and para.style.name:
                style_name = para.style.name
        except Exception:
            style_name = ""

        runs = self._collect_runs(para)
        eq = self._equation_text(para)
        if eq:
            runs.append({"text": f"[Equation: {eq}]"})

        text = "".join(r.get("text", "") for r in runs).strip()
        list_kind, list_level = self._list_info(para, numbering)

        blocks: list[dict[str, Any]] = []
        if text:
            heading_level = self._heading_level(style_name)
            if heading_level is not None and list_kind is None:
                blocks.append({"type": "heading", "level": heading_level, "runs": runs})
            elif self._is_code_style(style_name) and list_kind is None:
                blocks.append({
                    "type": "code",
                    "language": None,
                    "content": "".join(r.get("text", "") for r in runs),
                })
            elif style_name in _QUOTE_STYLES and list_kind is None:
                blocks.append({
                    "type": "blockquote",
                    "blocks": [{"order": 0, "type": "paragraph", "runs": runs}],
                })
            else:
                pb: dict[str, Any] = {"type": "paragraph", "runs": runs}
                if list_kind is not None:
                    pb["list_kind"] = list_kind
                    pb["list_level"] = list_level
                elif style_name and style_name not in _DEFAULT_STYLES:
                    pb["style"] = style_name
                blocks.append(pb)

        # Drawings (inline images + floating textboxes) trail the paragraph text.
        blocks.extend(self._drawing_blocks(para, doc, images, counter, warnings))
        return blocks

    def _collect_runs(self, para: Paragraph) -> list[dict[str, Any]]:
        """Build TextRuns preserving order across plain runs and hyperlinks."""
        out: list[dict[str, Any]] = []
        try:
            inner = list(para.iter_inner_content())
        except Exception:
            inner = list(para.runs)
        for item in inner:
            if isinstance(item, Hyperlink):
                url = getattr(item, "address", "") or getattr(item, "url", "") or ""
                for r in item.runs:
                    tr = self._run_to_textrun(r, hyperlink=url or None)
                    if tr is not None:
                        out.append(tr)
            elif isinstance(item, Run):
                tr = self._run_to_textrun(item)
                if tr is not None:
                    out.append(tr)
        return out

    def _run_to_textrun(self, run: Run, hyperlink: str | None = None) -> dict[str, Any] | None:
        text = run.text or ""
        if text == "" and hyperlink is None:
            return None
        tr: dict[str, Any] = {"text": text}
        if run.bold:
            tr["bold"] = True
        if run.italic:
            tr["italic"] = True
        if run.underline:
            tr["underline"] = True
        if hyperlink:
            tr["hyperlink"] = hyperlink
        return tr

    def _list_info(
        self, para: Paragraph, numbering: dict[str, dict[int, str]]
    ) -> tuple[str | None, int]:
        # Numbering lives either as direct paragraph formatting (pPr/numPr) or,
        # for named list styles ("List Bullet" / "List Number"), on the style.
        numPr = self._find_numpr(para._p)
        if numPr is None:
            try:
                style_el = para.style.element if para.style is not None else None
            except Exception:
                style_el = None
            if style_el is not None:
                numPr = self._find_numpr(style_el)
        if numPr is None:
            return None, 0

        numid_el = numPr.find(qn("w:numId"))
        numId = numid_el.get(qn("w:val")) if numid_el is not None else None
        if not numId or numId == "0":
            return None, 0
        ilvl = 0
        ilvl_el = numPr.find(qn("w:ilvl"))
        if ilvl_el is not None:
            try:
                ilvl = int(ilvl_el.get(qn("w:val")))
            except (TypeError, ValueError):
                ilvl = 0
        kind = numbering.get(numId, {}).get(ilvl, "bullet")
        return kind, ilvl

    def _find_numpr(self, el: Any) -> Any:
        pPr = el.find(qn("w:pPr"))
        return pPr.find(qn("w:numPr")) if pPr is not None else None

    def _heading_level(self, style_name: str) -> int | None:
        if not style_name:
            return None
        if style_name == "Title":
            return 1
        if style_name == "Subtitle":
            return 2
        m = _HEADING_RE.search(style_name)
        if m:
            return min(int(m.group(1)), 6)
        return None

    def _is_code_style(self, style_name: str) -> bool:
        return bool(style_name) and ("code" in style_name.lower() or style_name in _CODE_STYLES)

    def _equation_text(self, para: Paragraph) -> str:
        omaths = para._p.findall(f".//{{{_M_NS}}}oMath")
        if not omaths:
            return ""
        parts: list[str] = []
        for om in omaths:
            texts = om.findall(f".//{{{_M_NS}}}t")
            joined = "".join(t.text or "" for t in texts).strip()
            if joined:
                parts.append(joined)
        return " ".join(parts).strip()

    # ----------------- drawings (images + textboxes) -----------------

    def _drawing_blocks(
        self,
        para: Paragraph,
        doc: Any,
        images: dict[str, bytes],
        counter: dict[str, int],
        warnings: list[str],
    ) -> list[dict[str, Any]]:
        blocks: list[dict[str, Any]] = []
        for drawing in para._p.findall(".//" + qn("w:drawing")):
            txbx = drawing.find(".//" + qn("w:txbxContent"))
            if txbx is not None:
                inner = self._textbox_blocks(txbx, doc, images, counter, warnings)
                if inner:
                    blocks.append({"type": "textbox", "blocks": inner})
                continue
            relids = drawing.find(f".//{{{_DGM_NS}}}relIds")
            if relids is not None:
                sa = self._smartart_block(relids, warnings)
                if sa is not None:
                    blocks.append(sa)
                continue
            blip = drawing.find(".//" + qn("a:blip"))
            if blip is None:
                continue
            rId = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            if not rId:
                continue
            block = self._image_block(rId, drawing, doc, images, counter, warnings)
            if block is not None:
                blocks.append(block)
        return blocks

    def _image_block(
        self,
        rId: str,
        drawing: Any,
        doc: Any,
        images: dict[str, bytes],
        counter: dict[str, int],
        warnings: list[str],
    ) -> dict[str, Any] | None:
        try:
            blob = doc.part.related_parts[rId].blob
        except Exception as e:
            warnings.append(f"docx image rId={rId} load failed: {e}")
            return None
        normalized = normalize_to_png(blob)
        w, h = image_dimensions(normalized)
        counter["img"] += 1
        img_id = f"img_{counter['img']}"
        rel_path = f"raw_images/{img_id}.png"
        images[rel_path] = normalized

        alt = ""
        docpr = drawing.find(".//" + qn("wp:docPr"))
        if docpr is not None:
            alt = docpr.get("descr") or docpr.get("title") or docpr.get("name") or ""

        image_ref: dict[str, Any] = {
            "id": img_id,
            "relative_path": rel_path,
            "width_px": w,
            "height_px": h,
            "container": {"kind": "docx_section", "index": 1, "label": None},
            "neighbors": {
                "before_text": "", "after_text": "",
                "before_kind": None, "after_kind": None,
            },
        }
        if alt:
            image_ref["alt_text"] = alt
        return {"type": "image", "image": image_ref}

    def _textbox_blocks(
        self,
        txbx: Any,
        doc: Any,
        images: dict[str, bytes],
        counter: dict[str, int],
        warnings: list[str],
    ) -> list[dict[str, Any]]:
        blocks: list[dict[str, Any]] = []
        for child in txbx.iterchildren():
            if isinstance(child, CT_P):
                para = Paragraph(child, doc)
                new = self._paragraph_to_blocks(para, doc, {}, images, counter, warnings)
            elif isinstance(child, CT_Tbl):
                new = [self._table_to_block(Table(child, doc))]
            else:
                continue
            for b in new:
                b["order"] = len(blocks)
                blocks.append(b)
        return blocks

    # ----------------- smartart (dgm diagrams) -----------------

    def _load_diagrams(self, source: Path, warnings: list[str]) -> None:
        """Pre-load diagrams/data*.xml + layout*.xml and the document rels.

        SmartArt is a `<w:drawing>` whose graphicData references the diagram
        data/layout parts via `<dgm:relIds r:dm=... r:lo=...>`; python-docx
        exposes neither, so read them straight from the package zip.
        """
        self._dgm_data: dict[str, bytes] = {}
        self._dgm_layout: dict[str, bytes] = {}
        self._doc_rels: dict[str, str] = {}
        try:
            with zipfile.ZipFile(source) as z:
                names = z.namelist()
                for name in names:
                    if name.startswith("word/diagrams/data") and name.endswith(".xml"):
                        self._dgm_data[name] = z.read(name)
                    elif name.startswith("word/diagrams/layout") and name.endswith(".xml"):
                        self._dgm_layout[name] = z.read(name)
                if "word/_rels/document.xml.rels" in names:
                    # Rel targets are relative to the part (word/document.xml) → word/.
                    for rid, target in parse_rels(z.read("word/_rels/document.xml.rels")).items():
                        if target.startswith("/"):
                            self._doc_rels[rid] = target.lstrip("/")
                        else:
                            self._doc_rels[rid] = posixpath.normpath(
                                posixpath.join("word", target)
                            )
        except Exception as e:
            warnings.append(f"docx diagram preload failed: {e}")

    def _smartart_block(self, relids: Any, warnings: list[str]) -> dict[str, Any] | None:
        data_xml = None
        layout_xml = None
        dm_rid = relids.get(qn("r:dm"))
        if dm_rid and dm_rid in self._doc_rels:
            data_xml = self._dgm_data.get(self._doc_rels[dm_rid])
        lo_rid = relids.get(qn("r:lo"))
        if lo_rid and lo_rid in self._doc_rels:
            layout_xml = self._dgm_layout.get(self._doc_rels[lo_rid])
        # Fallback: single-diagram docs resolve fine without the rel hop.
        if data_xml is None and self._dgm_data:
            data_xml = next(iter(self._dgm_data.values()))
        if layout_xml is None and self._dgm_layout:
            layout_xml = next(iter(self._dgm_layout.values()))
        if data_xml is None:
            warnings.append("docx SmartArt data XML not found")
            return None
        try:
            nodes = parse_smartart_tree(data_xml)
        except Exception as e:
            warnings.append(f"docx SmartArt parse failed: {e}")
            return None
        layout_kind = detect_layout_kind(layout_xml) if layout_xml else "other"
        return {"type": "smartart", "layout_kind": layout_kind, "nodes": nodes}

    # ----------------- tables -----------------

    def _table_to_block(self, table: Table) -> dict[str, Any]:
        rows: list[list[str]] = []
        for row in table.rows:
            cells: list[str] = []
            prev_tc = None
            for cell in row.cells:
                tc = cell._tc
                # Horizontally merged cells repeat the same <w:tc>; blank the dups.
                cells.append("" if tc is prev_tc else cell.text.strip())
                prev_tc = tc
            rows.append(cells)
        return {"type": "table", "rows": rows}

    # ----------------- numbering -----------------

    def _build_numbering_map(
        self, doc: Any, warnings: list[str]
    ) -> dict[str, dict[int, str]]:
        """Map numId -> {ilvl: "bullet"|"numbered"} from the numbering part."""
        try:
            root = doc.part.numbering_part.element
        except Exception:
            return {}

        abstract_fmts: dict[str, dict[int, str]] = {}
        for absnum in root.findall(qn("w:abstractNum")):
            aid = absnum.get(qn("w:abstractNumId"))
            if aid is None:
                continue
            lvls: dict[int, str] = {}
            for lvl in absnum.findall(qn("w:lvl")):
                fmt_el = lvl.find(qn("w:numFmt"))
                fmt = fmt_el.get(qn("w:val")) if fmt_el is not None else None
                try:
                    ilvl = int(lvl.get(qn("w:ilvl")))
                except (TypeError, ValueError):
                    continue
                lvls[ilvl] = "bullet" if fmt == "bullet" else "numbered"
            abstract_fmts[aid] = lvls

        out: dict[str, dict[int, str]] = {}
        for num in root.findall(qn("w:num")):
            nid = num.get(qn("w:numId"))
            ref = num.find(qn("w:abstractNumId"))
            if nid is not None and ref is not None:
                out[nid] = abstract_fmts.get(ref.get(qn("w:val")), {})
        return out

    # ----------------- image neighbors -----------------

    def _wire_image_neighbors(self, blocks: list[dict[str, Any]]) -> None:
        for i, b in enumerate(blocks):
            if b.get("type") != "image":
                continue
            before_text = after_text = ""
            before_kind = after_kind = None
            if i > 0:
                prev = blocks[i - 1]
                before_kind = prev.get("type")
                before_text = self._block_text(prev)
            if i + 1 < len(blocks):
                nxt = blocks[i + 1]
                after_kind = nxt.get("type")
                after_text = self._block_text(nxt)
            b["image"]["neighbors"] = {
                "before_text": before_text[:200],
                "after_text": after_text[:200],
                "before_kind": before_kind,
                "after_kind": after_kind,
            }

    def _block_text(self, block: dict[str, Any]) -> str:
        t = block.get("type")
        if t in ("heading", "paragraph"):
            return "".join(r.get("text", "") for r in block.get("runs", []))
        if t == "table":
            return " | ".join(" ".join(str(c) for c in row) for row in block.get("rows", [])[:2])
        if t == "code":
            return block.get("content", "")[:200]
        if t == "blockquote":
            inner = block.get("blocks", [])
            return self._block_text(inner[0]) if inner else ""
        if t == "smartart":
            texts: list[str] = []
            stack = list(block.get("nodes", []))
            while stack and len(texts) < 8:
                node = stack.pop(0)
                if node.get("text"):
                    texts.append(node["text"])
                stack[:0] = node.get("children", [])
            return "SmartArt: " + ", ".join(texts)
        return ""
